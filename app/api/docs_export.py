from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from datetime import datetime, timedelta, timezone
from app.utils.timezone import get_now_vn
import os
import io
from docx import Document
from docx.shared import Inches
from app.extensions import db
from app.models.task import Task
from app.models.user import User
from app.models.whiteboard import Whiteboard, WhiteboardElement
from app.services.google_service import GoogleAPIService
import json

bp = Blueprint('docs_export', __name__)


@bp.route('/task/<int:task_id>/export/word', methods=['GET'])
@jwt_required()
def export_task_to_word(task_id):
    """Export task details to Word document"""
    user_id = get_jwt_identity()
    task = Task.query.filter_by(id=task_id, user_id=user_id).first_or_404()
    
    # Create Word document
    doc = Document()
    
    # Add title
    doc.add_heading(f'Task: {task.name}', 0)
    
    # Add metadata
    doc.add_heading('Task Details', level=1)
    doc.add_paragraph(f'ID: {task.id}')
    doc.add_paragraph(f'Created: {task.created_at.strftime("%Y-%m-%d %H:%M:%S") if task.created_at else "N/A"}')
    doc.add_paragraph(f'Deadline: {task.deadline.strftime("%Y-%m-%d %H:%M:%S") if task.deadline else "N/A"}')
    doc.add_paragraph(f'Progress: {task.state}%')
    doc.add_paragraph(f'Completed: {"Yes" if task.is_done else "No"}')
    doc.add_paragraph(f'Price: ${task.price}' if task.price else 'Price: Not set')
    
    # Add content
    if task.content:
        doc.add_heading('Description', level=1)
        doc.add_paragraph(task.content)
    
    # Add client information
    if task.client_num or task.client_mail:
        doc.add_heading('Client Information', level=1)
        if task.client_num:
            doc.add_paragraph(f'Phone: {task.client_num}')
        if task.client_mail:
            doc.add_paragraph(f'Email: {task.client_mail}')
    
    # Add notes
    if task.noted:
        doc.add_heading('Notes', level=1)
        doc.add_paragraph(task.noted)
    
    # Save to in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Return file for download
    from werkzeug.utils import secure_filename
    filename = secure_filename(f"task_{task_id}_{task.name.replace(' ', '_')}.docx")
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@bp.route('/task/<int:task_id>/export/gg-docs', methods=['POST'])
@jwt_required()
def export_task_to_gg_docs(task_id):
    """Export task details to Google Docs (returns URL)"""
    user_id = get_jwt_identity()
    task = Task.query.filter_by(id=task_id, user_id=user_id).first_or_404()
    
    # Create document content
    doc_content = f"""Task: {task.name}
ID: {task.id}
Created: {task.created_at.strftime("%Y-%m-%d %H:%M:%S") if task.created_at else "N/A"}
Deadline: {task.deadline.strftime("%Y-%m-%d %H:%M:%S") if task.deadline else "N/A"}
Progress: {task.state}%
Completed: {"Yes" if task.is_done else "No"}
Price: ${task.price if task.price else 0}

Description:
{task.content or "No description provided"}

Client Information:
Phone: {task.client_num or "N/A"}
Email: {task.client_mail or "N/A"}

Notes:
{task.noted or "No notes provided"}
"""
    
    # Create Google Doc
    google_service = GoogleAPIService()
    result = google_service.create_document(
        title=f"Task: {task.name}",
        content=doc_content
    )
    
    if 'error' in result:
        # If Google API is not configured, return the content for manual creation
        return jsonify({
            'message': 'Google Docs API not configured',
            'task_id': task.id,
            'task_name': task.name,
            'content': doc_content,
            'instructions': 'Set up Google API credentials in environment variables to enable automatic export'
        })
    
    return jsonify({
        'task_id': task.id,
        'task_name': task.name,
        'document_id': result['document_id'],
        'url': result['url'],
        'title': result['title']
    })


@bp.route('/task/<int:task_id>/export/gg-forms', methods=['POST'])
@jwt_required()
def create_gg_forms(task_id):
    """Create a Google Form based on task details"""
    user_id = get_jwt_identity()
    task = Task.query.filter_by(id=task_id, user_id=user_id).first_or_404()
    
    # In a real implementation, this would connect to Google Forms API
    # For now, we'll return a placeholder with instructions
    form_structure = {
        'title': f'Task: {task.name}',
        'description': f'Feedback form for task #{task.id}',
        'questions': [
            {
                'type': 'TEXT',
                'title': 'Task ID',
                'required': True,
                'default': str(task.id)
            },
            {
                'type': 'TEXT',
                'title': 'Task Name',
                'required': True,
                'default': task.name
            },
            {
                'type': 'PARAGRAPH',
                'title': 'Task Description',
                'default': task.content or 'No description'
            },
            {
                'type': 'SCALE',
                'title': 'Task Progress Rating (0-100)',
                'options': list(range(0, 101, 10))
            },
            {
                'type': 'CHECKBOX',
                'title': 'Task Status',
                'options': ['Completed', 'In Progress', 'Not Started']
            },
            {
                'type': 'TEXT',
                'title': 'Additional Comments',
                'required': False
            }
        ]
    }
    
    return jsonify({
        'message': 'Google Forms creation would be implemented here',
        'task_id': task.id,
        'form_structure': form_structure,
        'instructions': 'To implement this feature, configure Google Forms API with OAuth credentials'
    })


@bp.route('/whiteboards', methods=['POST'])
@jwt_required()
def create_whiteboard():
    """Create a new whiteboard for the user"""
    user_id = get_jwt_identity()
    data = request.get_json()
    
    whiteboard = Whiteboard(
        user_id=user_id,
        name=data.get('name', f'Whiteboard_{get_now_vn().strftime("%Y%m%d_%H%M%S")}'),
        description=data.get('description', '')
    )
    
    db.session.add(whiteboard)
    db.session.commit()
    
    return jsonify({
        'id': whiteboard.id,
        'name': whiteboard.name,
        'description': whiteboard.description,
        'created_at': whiteboard.created_at.isoformat() if whiteboard.created_at else None,
        'user_id': whiteboard.user_id
    }), 201


@bp.route('/whiteboard/<int:whiteboard_id>', methods=['GET'])
@jwt_required()
def get_whiteboard(whiteboard_id):
    """Get whiteboard details"""
    user_id = get_jwt_identity()
    
    whiteboard = Whiteboard.query.filter_by(id=whiteboard_id, user_id=user_id).first_or_404()
    
    return jsonify(whiteboard.to_dict())


@bp.route('/whiteboard/<int:whiteboard_id>/export', methods=['POST'])
@jwt_required()
def export_whiteboard(whiteboard_id):
    """Export whiteboard as image or document"""
    user_id = get_jwt_identity()
    data = request.get_json()
    
    whiteboard = Whiteboard.query.filter_by(id=whiteboard_id, user_id=user_id).first_or_404()
    
    export_format = data.get('format', 'png')  # png, jpg, pdf, svg
    
    # Placeholder implementation - in real implementation, would render whiteboard to image
    return jsonify({
        'whiteboard_id': whiteboard_id,
        'export_format': export_format,
        'message': f'Whiteboard export to {export_format} would be implemented with canvas rendering',
        'whiteboard_data': whiteboard.data
    })


@bp.route('/whiteboard/<int:whiteboard_id>/elements', methods=['GET'])
@jwt_required()
def get_whiteboard_elements(whiteboard_id):
    """Get all elements in a whiteboard"""
    user_id = get_jwt_identity()
    
    whiteboard = Whiteboard.query.filter_by(id=whiteboard_id, user_id=user_id).first_or_404()
    
    elements = WhiteboardElement.query.filter_by(whiteboard_id=whiteboard_id).all()
    
    return jsonify({
        'whiteboard_id': whiteboard_id,
        'elements': [element.to_dict() for element in elements],
        'count': len(elements)
    })


@bp.route('/whiteboard/<int:whiteboard_id>/elements', methods=['POST'])
@jwt_required()
def add_whiteboard_element(whiteboard_id):
    """Add an element to a whiteboard"""
    user_id = get_jwt_identity()
    data = request.get_json()
    
    whiteboard = Whiteboard.query.filter_by(id=whiteboard_id, user_id=user_id).first_or_404()
    
    element = WhiteboardElement(
        whiteboard_id=whiteboard_id,
        type=data['type'],
        content=data.get('content', {}),
        position_x=data.get('position_x', 0),
        position_y=data.get('position_y', 0),
        width=data.get('width', 100),
        height=data.get('height', 50)
    )
    
    db.session.add(element)
    db.session.commit()
    
    return jsonify(element.to_dict()), 201


@bp.route('/whiteboard/<int:whiteboard_id>/elements/<int:element_id>', methods=['PUT'])
@jwt_required()
def update_whiteboard_element(whiteboard_id, element_id):
    """Update a whiteboard element"""
    user_id = get_jwt_identity()
    data = request.get_json()
    
    whiteboard = Whiteboard.query.filter_by(id=whiteboard_id, user_id=user_id).first_or_404()
    element = WhiteboardElement.query.filter_by(
        id=element_id, 
        whiteboard_id=whiteboard_id
    ).first_or_404()
    
    if 'type' in data:
        element.type = data['type']
    if 'content' in data:
        element.content = data['content']
    if 'position_x' in data:
        element.position_x = data['position_x']
    if 'position_y' in data:
        element.position_y = data['position_y']
    if 'width' in data:
        element.width = data['width']
    if 'height' in data:
        element.height = data['height']
    
    db.session.commit()
    
    return jsonify(element.to_dict())


@bp.route('/whiteboard/<int:whiteboard_id>/elements/<int:element_id>', methods=['DELETE'])
@jwt_required()
def delete_whiteboard_element(whiteboard_id, element_id):
    """Delete a whiteboard element"""
    user_id = get_jwt_identity()
    
    whiteboard = Whiteboard.query.filter_by(id=whiteboard_id, user_id=user_id).first_or_404()
    element = WhiteboardElement.query.filter_by(
        id=element_id, 
        whiteboard_id=whiteboard_id
    ).first_or_404()
    
    db.session.delete(element)
    db.session.commit()
    
    return jsonify({'message': 'Element deleted successfully'})


@bp.route('/whiteboards', methods=['GET'])
@jwt_required()
def get_user_whiteboards():
    """Get all whiteboards for the current user"""
    user_id = get_jwt_identity()
    
    whiteboards = Whiteboard.query.filter_by(user_id=user_id).all()
    
    return jsonify({
        'whiteboards': [wb.to_dict() for wb in whiteboards],
        'count': len(whiteboards)
    })
